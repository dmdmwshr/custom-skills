import os
import sys
import json
import shutil

def convert_doc_to_docx(doc_path):
    import win32com.client
    # 为了防止原文件被前面崩溃的脚本锁死，我们先copy一个临时脱壳版本来读
    temp_doc = doc_path + ".temp.doc"
    shutil.copy(doc_path, temp_doc)
    
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    
    docx_path = doc_path + "x"
    if os.path.exists(docx_path):
        try: os.remove(docx_path)
        except: pass
        
    try:
        doc = word.Documents.Open(os.path.abspath(temp_doc), ReadOnly=True)
        # 16 = wdFormatDocumentDefault
        doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)
        doc.Close(False)
    except Exception as e:
        print(f"Failed to convert to docx: {e}")
    finally:
        word.Quit()
        try: os.remove(temp_doc)
        except: pass
        
    return docx_path

def extract_from_docx(docx_path):
    from docx import Document
    document = Document(docx_path)
    
    if not document.tables:
        return {}
        
    table = document.tables[0]
    
    grid = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        grid.append(row_data)

    record = {
        "doc_name": os.path.splitext(os.path.basename(docx_path))[0],
        "fields": {},
        "deductions": []
    }

    # 抽取核心 fields
    for r in grid:
        for c_idx, cell_text in enumerate(r):
            if "评查组织单位" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["评查组织单位"] = r[c_idx + 1]
            if "评查人" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["评查人"] = r[c_idx + 1]
            if "复核人" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["复核人"] = r[c_idx + 1]
            if "被检查单位" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["被检查单位"] = r[c_idx + 1]
            if "评查日期" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["评查日期"] = r[c_idx + 1]
            if "评查形式" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["评查形式"] = r[c_idx + 1]
            if "题名" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["题名"] = r[c_idx + 1]
            if "编号" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["编号"] = r[c_idx + 1]
            if "立卷人" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["立卷人"] = r[c_idx + 1]
            if "检查人" in cell_text and c_idx + 1 < len(r): 
                record["fields"]["检查人"] = r[c_idx + 1]
                
    # 抽取 deductions 与分数
    for r in grid:
        try:
            if len(r) >= 5 and r[0].isdigit():
                idx = int(r[0])
                item_name = r[1]
                res = r[2].strip()
                rev = r[3].strip()
                note = r[4].strip()
                
                if res or rev or note:
                    if idx <= 4 and ("主体" in item_name or "证据" in item_name or "法律法规" in item_name or "程序" in item_name):
                        deduction = {"type": "基本要素", "序号": idx}
                        if res: deduction["评查结果"] = res
                        if rev: deduction["复核情况"] = rev
                        if note: deduction["情况说明"] = note
                        if deduction not in record["deductions"]:
                            record["deductions"].append(deduction)
                    else:
                        deduction = {"type": "一般要素", "序号": idx}
                        if res: deduction["扣分情况"] = res
                        if rev: deduction["复核情况"] = rev
                        if note: deduction["扣分说明"] = note
                        if deduction not in record["deductions"]:
                            record["deductions"].append(deduction)
                        
            if "合计扣分" in r:
                for c_idx, txt in enumerate(r):
                    if "合计扣分" in txt and c_idx + 1 < len(r):
                        val = r[c_idx+1].strip()
                        if val:
                            record["fields"]["合计扣分"] = val
            
            if "实际得分" in r:
                for c_idx, txt in enumerate(r):
                    if "实际得分" in txt and c_idx + 1 < len(r):
                        val = r[c_idx+1].strip()
                        if val:
                            record["fields"]["实际得分"] = val
        except Exception as e:
            pass
            
    return record


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python reader.py <input.doc> <output.json>")
        sys.exit(1)
        
    doc_path = sys.argv[1]
    out_path = sys.argv[2]
    
    docx_tmp = convert_doc_to_docx(doc_path)
    if os.path.exists(docx_tmp):
        extracted = extract_from_docx(docx_tmp)
        with open(out_path, 'w', encoding='utf-8') as f:
            json.dump([extracted], f, ensure_ascii=False, indent=2)
        print(f"✅ Extracted data safely via DOCX and saved to: {out_path}")
        try: os.remove(docx_tmp)
        except: pass
    else:
        print("❌ Failed to create temporary docx")
