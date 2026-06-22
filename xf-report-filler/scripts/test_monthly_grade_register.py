import copy
import sys
import tempfile
import types
import unittest
from datetime import date
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from openpyxl import Workbook, load_workbook


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import monthly_grade_register as mgr


class MonthlyGradeRegisterMonthTests(unittest.TestCase):
    def test_resolve_score_month_prefers_filename(self):
        info = mgr.resolve_score_month(
            r"C:\workspace\5月\2026年5月通报.doc",
            default_year=2026,
        )
        self.assertEqual(info["source"], "filename")
        self.assertEqual((info["score_year"], info["score_month"]), (2026, 5))
        self.assertEqual(info["score_month_key"], "2026-05")

    def test_resolve_score_month_from_folder(self):
        info = mgr.resolve_score_month(
            r"C:\workspace\5月\通报.doc",
            default_year=2026,
        )
        self.assertEqual(info["source"], "folder")
        self.assertEqual((info["score_year"], info["score_month"]), (2026, 5))

    def test_resolve_score_month_falls_back_to_runtime_date(self):
        info = mgr.resolve_score_month(
            r"C:\workspace\通报.doc",
            fallback_date=date(2026, 5, 21),
        )
        self.assertEqual(info["source"], "runtime")
        self.assertEqual((info["score_year"], info["score_month"]), (2026, 5))

    def test_resolve_score_month_cross_year(self):
        info = mgr.resolve_score_month(
            r"C:\workspace\1月\2026年1月通报.doc",
            default_year=2026,
        )
        self.assertEqual((info["score_year"], info["score_month"]), (2026, 1))
        self.assertEqual(info["score_month_key"], "2026-01")

    def test_build_generated_history_record_uses_score_month(self):
        record = mgr.build_generated_history_record(
            2026,
            5,
            r"C:\workspace\5月\2026年5月通报.doc",
            "（五）联网监测",
            [],
            {},
        )
        self.assertEqual(record["month_key"], "2026-05")
        self.assertEqual((record["year"], record["month"]), (2026, 5))

    def test_migrate_monitor_history_records_keeps_existing_months(self):
        history = {
            "version": 1,
            "records": [
                {
                    "year": 2026,
                    "month": 3,
                    "month_key": "2026-03",
                    "report_file": r"C:\workspace\3月\2026年3月通报.doc",
                    "selected_cases": [],
                    "source": "scanned_existing",
                },
                {
                    "year": 2026,
                    "month": 4,
                    "month_key": "2026-04",
                    "report_file": r"C:\workspace\4月\2026年4月通报.doc",
                    "selected_cases": [],
                    "source": "scanned_existing",
                },
                {
                    "year": 2026,
                    "month": 5,
                    "month_key": "2026-05",
                    "report_file": r"C:\workspace\5月\2026年5月通报.doc",
                    "selected_cases": [],
                    "source": "generated",
                },
            ],
        }
        migrated, actions = mgr.migrate_monitor_history_records(history)
        self.assertEqual([item["month_key"] for item in migrated["records"]], ["2026-03", "2026-04", "2026-05"])
        self.assertEqual(actions, [])

    def test_migrate_monitor_history_records_raises_on_conflict(self):
        history = {
            "version": 1,
            "records": [
                {
                    "year": 2026,
                    "month": 5,
                    "month_key": "2026-05",
                    "report_file": r"C:\workspace\5月\2026年5月通报.doc",
                    "selected_cases": [],
                    "source": "generated",
                },
                {
                    "year": 2026,
                    "month": 5,
                    "month_key": "2026-05-copy",
                    "report_file": r"C:\workspace\5月\另一份2026年5月通报.doc",
                    "selected_cases": [],
                    "source": "generated",
                },
            ],
        }
        with self.assertRaises(mgr.HistoryMonthConflict):
            mgr.migrate_monitor_history_records(copy.deepcopy(history))

    def test_find_base_info_prefers_root_normal_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            network = base / "2026年5月联网监测基础信息考评明细表"
            network.mkdir()
            root_file = base / "2026年5月基础信息考评截图（不发）.xls"
            nested_file = network / "5月基础信息考评截图.xls"
            root_file.write_text("root", encoding="utf-8")
            nested_file.write_text("nested", encoding="utf-8")
            self.assertEqual(mgr.find_base_info(base, network), root_file)

    def test_find_base_info_keeps_pending_prefix_compatibility(self):
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            network = base / "2026年5月联网监测基础信息考评明细表"
            network.mkdir()
            root_file = base / "【待补】5月基础信息考评截图（不发）.xls"
            root_file.write_text("root", encoding="utf-8")
            self.assertEqual(mgr.find_base_info(base, network), root_file)

    def test_find_base_info_falls_back_to_network_dir(self):
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            network = base / "2026年5月联网监测基础信息考评明细表"
            network.mkdir()
            nested_file = network / "5月基础信息考评截图.xls"
            nested_file.write_text("nested", encoding="utf-8")
            self.assertEqual(mgr.find_base_info(base, network), nested_file)

    def test_parse_monitor_note_rejects_pseudo_contact(self):
        person, issues = mgr.parse_monitor_note("3202U01557  消防机构联系人、cad、pdf")
        self.assertEqual(person, "")
        self.assertIn("联系人", issues)

    def test_normalize_monitor_issues_uses_public_cad_pdf_names(self):
        text = mgr.normalize_monitor_issues("cad、CAD、缺cad、缺CAD、CAD图、点位图、pdf、PDF、缺pdf、缺PDF、火灾防控图")
        self.assertNotIn("CAD", text)
        self.assertNotIn("PDF", text)
        self.assertIn("缺点位图", text)
        self.assertIn("缺火灾防控图", text)

    def test_monitor_office_text_uses_public_cad_pdf_names(self):
        text = mgr.monitor_office_text(
            "梁溪",
            [
                {
                    "short": "梁溪",
                    "单位": "测试单位",
                    "issues_raw": "cad、pdf",
                }
            ],
        )
        self.assertEqual(text, "1、测试单位缺点位图、缺火灾防控图")

    def test_blank_product_template_item_is_skipped(self):
        self.assertTrue(mgr.is_blank_product_template_item("（）", "扣分："))
        self.assertTrue(mgr.is_blank_product_template_item(mgr.clean_error_line("3、（）"), "扣分："))
        self.assertFalse(mgr.is_blank_product_template_item("责令限期改正通知书填写不规范", "扣分：3.2 0.1分"))

    def test_product_public_description_uses_text_before_parentheses(self):
        cases = {
            "卷内应有的某文书缺失（缺少责令限期改正通知书审批表、证据保全清单）": "卷内应有的某文书缺失",
            "复查现场照片不齐全（缺少门头）": "复查现场照片不齐全",
            "缺授权委托书（缺授权委托书）": "缺授权委托书",
            "消防产品监督检查记录记录不规范（备注栏未填写）": "消防产品监督检查记录记录不规范",
            "责令限期改正通知书填写不规范 ps：具体问题不能进通报": "责令限期改正通知书填写不规范",
        }
        for raw, expected in cases.items():
            with self.subTest(raw=raw):
                self.assertEqual(mgr.public_product_description(raw), expected)
                self.assertEqual(mgr.normalize_broad_description(raw), expected)

    def test_validate_person_matches_allows_missing_monitor_contact(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "个人执法统计表模板.xlsx"
            workbook = Workbook()
            ws = workbook.active
            ws.cell(6, 1).value = "宜兴大队"
            ws.cell(6, 2).value = "张三"
            workbook.save(template)

            mgr.validate_person_matches(
                template,
                [],
                [
                    {
                        "short": "宜兴",
                        "大队": "宜兴大队",
                        "联系人": "",
                        "单位": "测试单位",
                        "note": "3202U01557  消防机构联系人、cad、pdf",
                    }
                ],
            )

    def test_product_detail_leak_guard_detects_parenthetical_detail(self):
        clean_record = {
            "大队": "梁溪大队",
            "short": "梁溪",
            "题名": "测试案卷",
            "立卷人": "张三",
            "no_case": False,
            "errors": ["责令限期改正通知书填写不规范（具体问题不能进通报）"],
            "archive_errors": ["责令限期改正通知书填写不规范"],
        }
        self.assertEqual(mgr.product_detail_leak_issues_for_records([clean_record]), [])
        leaked = copy.deepcopy(clean_record)
        leaked["archive_errors"] = ["责令限期改正通知书填写不规范（具体问题不能进通报）"]
        issues = mgr.product_detail_leak_issues_for_records([leaked])
        self.assertEqual(issues[0]["type"], "product_detail_leak")
        self.assertEqual(issues[0]["fragment"], "具体问题不能进通报")

    def test_product_public_description_guard_detects_semantic_detail_rewrite(self):
        record = {
            "大队": "滨湖大队",
            "short": "滨湖",
            "题名": "测试案卷",
            "立卷人": "张三",
            "no_case": False,
            "errors": ["卷内应有的某文书缺失（缺少责令限期改正通知书审批表、证据保全清单）"],
            "archive_errors": ["责令限期改正通知书审批材料缺失"],
        }
        issues = mgr.product_detail_leak_issues_for_records([record])
        mismatch = [item for item in issues if item["type"] == "product_public_description_mismatch"]
        self.assertEqual(mismatch[0]["expected"], ["卷内应有的某文书缺失"])
        self.assertEqual(mismatch[0]["actual"], ["责令限期改正通知书审批材料缺失"])

    def test_yellow_product_register_issue_is_private_and_not_scored(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "2026年5月产品巡查底册（不发）.docx"
            doc = Document()
            doc.add_paragraph("大队：宜兴大队")
            doc.add_paragraph("题名：测试案卷")
            doc.add_paragraph("编号：3202")
            doc.add_paragraph("立卷人：张三")
            doc.add_paragraph("检查人：李四")
            doc.add_paragraph("错误")
            doc.add_paragraph("1、缺授权委托书（缺授权委托书）")
            doc.add_paragraph("扣分：3.11 0.2分")
            private_issue = doc.add_paragraph()
            private_issue.add_run("2、卷内应有的某文书缺失（私账问题）").font.highlight_color = WD_COLOR_INDEX.YELLOW
            private_deduction = doc.add_paragraph()
            private_deduction.add_run("扣分：3.13 a 0.5分").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.save(path)

            records = mgr.parse_product_register(path)

            self.assertEqual(len(records), 1)
            record = records[0]
            self.assertEqual(record["score"], 9.8)
            self.assertEqual(record["errors"], ["缺授权委托书（缺授权委托书）"])
            self.assertEqual(record["archive_errors"], ["缺授权委托书"])
            self.assertEqual(len(record["deductions"]), 1)
            self.assertEqual(record["deductions"][0]["value"], 0.2)
            self.assertEqual(record["ignored_yellow_errors"][0]["description"], "卷内应有的某文书缺失（私账问题）")
            self.assertNotIn("私账问题", mgr.product_office_text(record))

    def test_write_personal_stats_marks_product_person_mismatch(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "个人执法统计表模板.xlsx"
            output = Path(tmp) / "2026年5月个人执法统计表.xlsx"
            workbook = Workbook()
            ws = workbook.active
            ws.cell(6, 1).value = "梁溪大队"
            ws.cell(6, 2).value = "张三"
            workbook.save(template)

            warnings = mgr.write_personal_stats(
                template,
                output,
                [
                    {
                        "short": "梁溪",
                        "大队": "梁溪大队",
                        "立卷人": "李四",
                        "题名": "测试案卷",
                        "score": 9.5,
                        "no_case": False,
                    }
                ],
                [],
                5,
                force=True,
            )
            wb = load_workbook(output)
            ws = wb.active
            self.assertIn("待核对", ws.cell(6, 27).value)
            self.assertIn("李四", ws.cell(6, 27).value)
            self.assertEqual(ws.cell(6, 27).fill.fgColor.rgb, "FFFF0000")
            self.assertIsNotNone(ws.cell(6, 27).comment)
            self.assertTrue(any("未找到产品立卷人" in item for item in warnings))

    def test_write_product_summary_formats_product_scores_one_decimal(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "产品监督成绩总表模板.xlsx"
            output = Path(tmp) / "产品监督成绩总表.xlsx"
            workbook = Workbook()
            ws = workbook.active
            ws.cell(2, 1).value = "梁溪大队"
            ws.cell(3, 1).value = "锡山大队"
            workbook.save(template)

            mgr.write_product_summary(
                template,
                output,
                [
                    {"short": "梁溪", "score": 10.0},
                    {"short": "锡山", "score": 9.5},
                ],
                force=True,
            )

            wb = load_workbook(output)
            ws = wb.active
            self.assertEqual(ws.cell(2, 2).value, 10)
            self.assertEqual(ws.cell(2, 2).number_format, mgr.PRODUCT_SCORE_NUMBER_FORMAT)
            self.assertEqual(ws.cell(3, 2).value, 9.5)
            self.assertEqual(ws.cell(3, 2).number_format, mgr.PRODUCT_SCORE_NUMBER_FORMAT)

    def test_write_personal_stats_formats_product_and_monitor_scores(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "个人执法统计表模板.xlsx"
            output = Path(tmp) / "2026年5月个人执法统计表.xlsx"
            workbook = Workbook()
            ws = workbook.active
            ws.cell(6, 1).value = "梁溪大队"
            ws.cell(6, 2).value = "张三"
            workbook.save(template)

            warnings = mgr.write_personal_stats(
                template,
                output,
                [
                    {
                        "short": "梁溪",
                        "大队": "梁溪大队",
                        "立卷人": "张三",
                        "题名": "测试案卷",
                        "score": 10.0,
                        "no_case": False,
                    }
                ],
                [
                    {
                        "short": "梁溪",
                        "大队": "梁溪大队",
                        "联系人": "张三",
                        "单位": "测试单位",
                        "score": 9.7,
                    }
                ],
                5,
                force=True,
            )

            wb = load_workbook(output)
            ws = wb.active
            self.assertEqual(warnings, [])
            self.assertEqual(ws.cell(6, 27).value, 10)
            self.assertEqual(ws.cell(6, 27).number_format, mgr.PRODUCT_SCORE_NUMBER_FORMAT)
            self.assertEqual(ws.cell(6, 28).value, 9.7)
            self.assertEqual(ws.cell(6, 28).number_format, mgr.MONITOR_SCORE_NUMBER_FORMAT)

    def test_write_case_scores_formats_product_and_monitor_scores(self):
        class FakeCell:
            def __init__(self, value=None):
                self.Value = value
                self.NumberFormat = None

        class FakeSheet:
            def __init__(self):
                self.cells = {(3, 1): FakeCell("梁溪大队")}

            def Cells(self, row, col):
                return self.cells.setdefault((row, col), FakeCell())

        class FakeWorkbook:
            def __init__(self, sheet):
                self.sheet = sheet
                self.saved = False
                self.closed = False

            def Worksheets(self, index):
                self.index = index
                return self.sheet

            def Save(self):
                self.saved = True

            def Close(self, save_changes):
                self.closed = True

        class FakeWorkbooks:
            def __init__(self, workbook):
                self.workbook = workbook
                self.opened_path = None

            def Open(self, path):
                self.opened_path = path
                return self.workbook

        class FakeExcel:
            def __init__(self, workbook):
                self.Workbooks = FakeWorkbooks(workbook)
                self.Visible = True
                self.DisplayAlerts = True
                self.quit = False

            def Quit(self):
                self.quit = True

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "个案成绩模板.xls"
            output = Path(tmp) / "个案成绩.xls"
            template.write_bytes(b"fake-xls")
            sheet = FakeSheet()
            workbook = FakeWorkbook(sheet)
            excel = FakeExcel(workbook)
            win32com = types.ModuleType("win32com")
            client = types.ModuleType("win32com.client")
            client.DispatchEx = lambda name: excel
            win32com.client = client

            with patch.dict(sys.modules, {"win32com": win32com, "win32com.client": client}):
                mgr.write_case_scores(
                    template,
                    output,
                    [{"short": "梁溪", "score": 10.0}],
                    {"梁溪": {"avg": 9.7}},
                    force=True,
                )

            self.assertTrue(workbook.saved)
            self.assertTrue(workbook.closed)
            self.assertEqual(sheet.Cells(3, 2).Value, 10.0)
            self.assertEqual(sheet.Cells(3, 2).NumberFormat, mgr.PRODUCT_SCORE_NUMBER_FORMAT)
            self.assertEqual(sheet.Cells(3, 11).Value, 9.7)
            self.assertEqual(sheet.Cells(3, 11).NumberFormat, mgr.MONITOR_SCORE_NUMBER_FORMAT)

    def test_read_monitor_scores_recomputes_avg_over_10(self):
        class FakeCell:
            def __init__(self, value):
                self.value = value

        class FakeSheet:
            title = "Sheet1"
            max_row = 3

            def __init__(self):
                self.values = {(3, 1): "惠山"}
                for col in range(2, 12):
                    self.values[(3, col)] = 9.7
                self.values[(3, 12)] = 97.2

            def cell(self, row, col):
                return FakeCell(self.values.get((row, col)))

        class FakeWorkbook:
            active = FakeSheet()

        with patch.object(mgr, "load_xls_as_workbook", return_value=FakeWorkbook()):
            scores = mgr.read_monitor_scores(Path("联网监测统计表.xls"))

        self.assertEqual(scores["惠山"]["avg"], 9.7)
        self.assertEqual(mgr.read_monitor_scores.last_warnings[0]["raw_avg"], 97.2)

    def test_default_output_files_do_not_include_score_office_record(self):
        outputs = mgr.build_output_files(Path(r"C:\workspace\6月通报\5月巡查"), 2026, 5)
        self.assertNotIn("office_record", outputs)
        self.assertEqual(outputs["monthly_report"], r"C:\workspace\6月通报\5月巡查\2026年5月通报.doc")

    def test_legacy_output_files_include_score_office_record_when_enabled(self):
        outputs = mgr.build_output_files(
            Path(r"C:\workspace\6月通报\5月巡查"),
            2026,
            5,
            include_score_office_record=True,
        )
        self.assertEqual(outputs["office_record"], r"C:\workspace\6月通报\5月巡查\2026年5月科室月考核情况记录表.xlsx")


if __name__ == "__main__":
    unittest.main()
