import argparse
import hashlib
import json
import os
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
SKILL_TEMPLATE_DIR = SKILL_DIR / "resources" / "monthly_templates"

DEFAULT_WORK_ROOT = Path(r"E:\文件夹\1、工作\2、产品，科技，联网检测\1、产品监督")
DEFAULT_MAY_DIR = DEFAULT_WORK_ROOT / "26年" / "5月通报"
DEFAULT_JUNE_MONTH_DIR = DEFAULT_WORK_ROOT / "26年" / "6月通报" / "5月巡查"
DEFAULT_BULLETIN_DIR = DEFAULT_WORK_ROOT / "26年" / "6月通报"
DEFAULT_TEMPLATE_DIR = DEFAULT_WORK_ROOT / "模板文件"

PENDING_PREFIX = "【待补】"

MAY_ROOT_RENAMES = {
    "5月科室月考核情况记录表.xlsx": "2026年5月科室月考核情况记录表.xlsx",
    "2026年5月工作检查-科室月考核情况记录表.xlsx": "2026年5月科室月考核情况记录表.xlsx",
    "5月消防产品监督统计表.xls": "2026年5月消防产品监督统计表.xls",
    "2026年5月工作检查-消防产品监督统计表.xls": "2026年5月消防产品监督统计表.xls",
    "5月重点工作完成情况上报表（应急通信与消防科技）.xls": "2026年5月重点工作完成情况上报表（应急通信与消防科技）.xls",
    "2026年5月工作检查-重点工作完成情况上报表（应急通信与消防科技）.xls": "2026年5月重点工作完成情况上报表（应急通信与消防科技）.xls",
}

MAY_PATROL_ROOT_RENAMES = {
    "（4月）产品巡查底册（不发）.docx": "2026年4月产品巡查底册（不发）.docx",
    "4月基础信息考评截图（不发）.xls": "2026年4月基础信息考评截图（不发）.xls",
    "个人执法统计表202604.xlsx": "2026年4月个人执法统计表.xlsx",
    "消防监督管理系统消防执法质量（4月个案成绩）.xls": "2026年4月消防监督管理系统消防执法质量（个案成绩）.xls",
}

TEMPLATE_RENAMES = {
    "（模板）产品巡查底册.docx": "01_产品巡查底册模板.docx",
    "（模板）消防产品档案质量明细表.doc": "02_消防产品档案质量明细表模板.doc",
    "（模板）产品监督成绩总表.xlsx": "03_产品监督成绩总表模板.xlsx",
    "（模板）个人执法统计表202601.xlsx": "04_个人执法统计表模板.xlsx",
    "（模板）科室月考核情况记录表.xlsx": "05_科室月考核情况记录表模板.xlsx",
    "(模板-成绩汇总)消防监督管理系统消防执法质量（个案成绩）.xls": "06_消防执法质量个案成绩模板.xls",
    "(样例)xxxx年x月通报.doc": "07_月度通报模板.doc",
    "（空表）每月消防产品监督统计表.xls": "08_每月消防产品监督统计表空表模板.xls",
    "(模板)X月消防产品监督统计表.xls": "09_消防产品监督统计表工作检查模板.xls",
    "（样例数据）消防产品档案质量明细表.doc": "90_消防产品档案质量明细表样例数据.doc",
}

MAY_TEMPLATE_SOURCES = {
    (
        "2026年5月重点工作完成情况上报表（应急通信与消防科技）.xls",
        "2026年5月工作检查-重点工作完成情况上报表（应急通信与消防科技）.xls",
        "5月重点工作完成情况上报表（应急通信与消防科技）.xls",
    ): "10_重点工作完成情况上报表（应急通信与消防科技）模板.xls",
}

NUMBERED_TEMPLATE_NAMES = [
    "01_产品巡查底册模板.docx",
    "02_消防产品档案质量明细表模板.doc",
    "03_产品监督成绩总表模板.xlsx",
    "04_个人执法统计表模板.xlsx",
    "05_科室月考核情况记录表模板.xlsx",
    "06_消防执法质量个案成绩模板.xls",
    "07_月度通报模板.doc",
    "08_每月消防产品监督统计表空表模板.xls",
    "09_消防产品监督统计表工作检查模板.xls",
    "10_重点工作完成情况上报表（应急通信与消防科技）模板.xls",
    "90_消防产品档案质量明细表样例数据.doc",
]

NUMBERED_TEMPLATE_CANDIDATES = {
    name: [name] for name in NUMBERED_TEMPLATE_NAMES
}
for old_name, new_name in TEMPLATE_RENAMES.items():
    NUMBERED_TEMPLATE_CANDIDATES.setdefault(new_name, [new_name]).append(old_name)

BULLETIN_ROOT_TEMPLATE_MAP = [
    {
        "library": "05_科室月考核情况记录表模板.xlsx",
        "skeleton": "X月科室月考核情况记录表.xlsx",
        "target": "{year}年{month}月科室月考核情况记录表.xlsx",
    },
    {
        "library": "09_消防产品监督统计表工作检查模板.xls",
        "skeleton": "X月消防产品监督统计表.xls",
        "target": "{year}年{month}月消防产品监督统计表.xls",
    },
    {
        "library": "10_重点工作完成情况上报表（应急通信与消防科技）模板.xls",
        "skeleton": "X月重点工作完成情况上报表（应急通信与消防科技）.xls",
        "target": "{year}年{month}月重点工作完成情况上报表（应急通信与消防科技）.xls",
    },
]

JUNE_TEMPLATE_COPY_NAMES = set(TEMPLATE_RENAMES) | {
    "（模板）产品监督成绩总表.xlsx",
    "（模板）个人执法统计表202601.xlsx",
    "（模板）科室月考核情况记录表.xlsx",
    "（模板）消防产品档案质量明细表.doc",
    "（空表）每月消防产品监督统计表.xls",
    "（样例数据）消防产品档案质量明细表.doc",
}

PRODUCT_REGISTER_PARAGRAPHS = [67, 68, 109, 110]

PRODUCT_REGISTER_NORMAL_NAME = "2026年5月产品巡查底册（不发）.docx"
BASE_INFO_NORMAL_NAME = "2026年5月基础信息考评截图（不发）.xls"


def sha256_file(path):
    digest = hashlib.sha256()
    with Path(path).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def is_under(path, root):
    path = Path(path).resolve()
    root = Path(root).resolve()
    return path == root or root in path.parents


def add_blocker(blockers, message, **extra):
    item = {"message": message}
    item.update({k: str(v) for k, v in extra.items()})
    blockers.append(item)


def move_or_rename(src, dst, actions, blockers, apply, root, kind):
    src = Path(src)
    dst = Path(dst)
    if not src.exists():
        return
    if not is_under(src, root) or not is_under(dst, root):
        add_blocker(blockers, "路径越界，已拒绝移动", src=src, dst=dst, root=root)
        return
    if dst.exists():
        if sha256_file(src) == sha256_file(dst):
            actions.append({"kind": kind, "status": "skip_same_hash_target_exists", "src": str(src), "dst": str(dst)})
            return
        add_blocker(blockers, "目标已存在且内容不同", src=src, dst=dst)
        return
    actions.append({"kind": kind, "status": "planned" if not apply else "done", "src": str(src), "dst": str(dst)})
    if apply:
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(dst))


def copy_if_missing(src, dst, actions, blockers, apply, root, kind):
    src = Path(src)
    dst = Path(dst)
    if not src.exists():
        add_blocker(blockers, "源文件不存在，无法复制", src=src, dst=dst)
        return
    if not is_under(src, root) or not is_under(dst, root):
        add_blocker(blockers, "路径越界，已拒绝复制", src=src, dst=dst, root=root)
        return
    if dst.exists():
        if sha256_file(src) == sha256_file(dst):
            actions.append({"kind": kind, "status": "skip_same_hash_target_exists", "src": str(src), "dst": str(dst)})
            return
        add_blocker(blockers, "目标已存在且内容不同", src=src, dst=dst)
        return
    actions.append({"kind": kind, "status": "planned" if not apply else "done", "src": str(src), "dst": str(dst)})
    if apply:
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)


def delete_file_if_exists(path, actions, blockers, apply, root, kind):
    path = Path(path)
    if not path.exists():
        return
    if not path.is_file():
        add_blocker(blockers, "目标不是文件，拒绝删除", path=path)
        return
    if not is_under(path, root):
        add_blocker(blockers, "路径越界，已拒绝删除", path=path, root=root)
        return
    digest = sha256_file(path)
    actions.append({"kind": kind, "status": "planned" if not apply else "done", "path": str(path), "sha256": digest})
    if apply:
        path.unlink()


def ensure_dir(path, actions, apply, kind):
    path = Path(path)
    if path.exists():
        actions.append({"kind": kind, "status": "skip_exists", "path": str(path)})
        return
    actions.append({"kind": kind, "status": "planned" if not apply else "done", "path": str(path)})
    if apply:
        path.mkdir(parents=True, exist_ok=True)


def common_root(*paths):
    resolved = [str(Path(path).resolve()) for path in paths]
    return Path(os.path.commonpath(resolved))


def numbered_template_dir(template_dir):
    return Path(template_dir) / "_编号模板库"


def skeleton_template_dir(template_dir):
    return Path(template_dir) / "X月通报"


def find_first_existing(paths):
    for path in paths:
        path = Path(path)
        if path.exists():
            return path
    return None


def template_source_for(template_dir, filename):
    template_dir = Path(template_dir)
    library_dir = numbered_template_dir(template_dir)
    candidates = [library_dir / filename]
    for candidate_name in NUMBERED_TEMPLATE_CANDIDATES.get(filename, [filename]):
        candidates.append(template_dir / candidate_name)
    return find_first_existing(candidates) or candidates[0]


def skeleton_source_for(template_dir, item):
    template_dir = Path(template_dir)
    skeleton = skeleton_template_dir(template_dir) / item["skeleton"]
    if skeleton.exists():
        return skeleton
    return template_source_for(template_dir, item["library"])


def organize_template_model(template_dir, actions, blockers, apply):
    template_dir = Path(template_dir)
    library_dir = numbered_template_dir(template_dir)
    skeleton_dir = skeleton_template_dir(template_dir)

    ensure_dir(library_dir, actions, apply, "ensure_numbered_template_library")
    ensure_dir(skeleton_dir, actions, apply, "ensure_bulletin_skeleton_dir")
    ensure_dir(skeleton_dir / "上月巡查", actions, apply, "ensure_bulletin_skeleton_patrol_dir")
    archive_office_locks(template_dir, actions, blockers, apply)

    for filename in NUMBERED_TEMPLATE_NAMES:
        if (library_dir / filename).exists():
            actions.append({"kind": "numbered_template_present", "status": "skip_exists", "path": str(library_dir / filename)})
            continue
        candidates = [template_dir / name for name in NUMBERED_TEMPLATE_CANDIDATES.get(filename, [filename])]
        source = find_first_existing(candidates)
        if source is None:
            add_blocker(blockers, "编号模板源文件不存在", target=library_dir / filename)
            continue
        move_or_rename(
            source,
            library_dir / filename,
            actions,
            blockers,
            apply,
            template_dir,
            "move_template_to_numbered_library",
        )

    for item in BULLETIN_ROOT_TEMPLATE_MAP:
        source = template_source_for(template_dir, item["library"])
        copy_if_missing(
            source,
            skeleton_dir / item["skeleton"],
            actions,
            blockers,
            apply,
            template_dir,
            "copy_bulletin_skeleton_template",
        )


def normalize_score_source_names(score_dir, score_year, score_month, actions, blockers, apply):
    score_dir = Path(score_dir)
    product_normal_name = f"{score_year}年{score_month}月产品巡查底册（不发）.docx"
    base_info_normal_name = f"{score_year}年{score_month}月基础信息考评截图（不发）.xls"
    network_dir = score_dir / f"{score_year}年{score_month}月联网监测基础信息考评明细表"

    product_target = score_dir / product_normal_name
    product_old_names = [
        f"{PENDING_PREFIX}（{score_month}月）产品巡查底册（不发）.docx",
        f"（{score_month}月）产品巡查底册（不发）.docx",
        f"（{score_month}月）产品巡查底册.docx",
        f"{score_month}月产品巡查底册（不发）.docx",
    ]
    for old_name in product_old_names:
        move_or_rename(score_dir / old_name, product_target, actions, blockers, apply, score_dir, "score_product_register_rename")

    base_info_target = score_dir / base_info_normal_name
    base_info_old_paths = [
        score_dir / f"{PENDING_PREFIX}{score_month}月基础信息考评截图（不发）.xls",
        score_dir / f"{score_month}月基础信息考评截图（不发）.xls",
        network_dir / f"{score_month}月基础信息考评截图.xls",
    ]
    for old_path in base_info_old_paths:
        move_or_rename(old_path, base_info_target, actions, blockers, apply, score_dir, "score_base_info_move")


def instantiate_bulletin_dir(bulletin_dir, bulletin_year, bulletin_month, score_year, score_month, template_dir, actions, blockers, apply):
    bulletin_dir = Path(bulletin_dir)
    template_dir = Path(template_dir)
    operation_root = common_root(bulletin_dir, template_dir)
    score_dir = bulletin_dir / f"{score_month}月巡查"

    ensure_dir(bulletin_dir, actions, apply, "ensure_bulletin_dir")
    ensure_dir(score_dir, actions, apply, "ensure_score_patrol_dir")

    for item in BULLETIN_ROOT_TEMPLATE_MAP:
        source = skeleton_source_for(template_dir, item)
        target_name = item["target"].format(year=bulletin_year, month=bulletin_month)
        copy_if_missing(
            source,
            bulletin_dir / target_name,
            actions,
            blockers,
            apply,
            operation_root,
            "copy_bulletin_root_file",
        )

    wrong_score_office = score_dir / f"{score_year}年{score_month}月科室月考核情况记录表.xlsx"
    delete_file_if_exists(
        wrong_score_office,
        actions,
        blockers,
        apply,
        bulletin_dir,
        "delete_wrong_score_office_record",
    )

    archive_june_template_copies(score_dir, numbered_template_dir(template_dir), actions, blockers, apply)
    normalize_score_source_names(score_dir, score_year, score_month, actions, blockers, apply)


def collect_known_template_hashes(template_dir):
    hashes = {}
    for directory in [Path(template_dir), SKILL_TEMPLATE_DIR]:
        if not directory.exists():
            continue
        for file_path in directory.glob("*"):
            if file_path.is_file() and not file_path.name.startswith("~$"):
                hashes.setdefault(sha256_file(file_path), []).append(str(file_path))
    return hashes


def archive_office_locks(base_dir, actions, blockers, apply):
    archive_dir = Path(base_dir) / "_Office临时锁文件归档"
    for lock_file in Path(base_dir).glob("~$*"):
        move_or_rename(lock_file, archive_dir / lock_file.name, actions, blockers, apply, base_dir, "archive_office_lock")


def organize_may_dir(may_dir, actions, blockers, apply):
    may_dir = Path(may_dir)
    for old_name, new_name in MAY_ROOT_RENAMES.items():
        move_or_rename(may_dir / old_name, may_dir / new_name, actions, blockers, apply, may_dir, "may_root_rename")

    patrol_dir = may_dir / "4月巡查"
    for old_name, new_name in MAY_PATROL_ROOT_RENAMES.items():
        move_or_rename(patrol_dir / old_name, patrol_dir / new_name, actions, blockers, apply, may_dir, "may_patrol_root_rename")

    product_dir = may_dir / "4月巡查" / "2026年4月消防产品监督成绩"
    if product_dir.exists():
        for file_path in product_dir.glob("*（5月产品监督档案）.doc"):
            move_or_rename(
                file_path,
                file_path.with_name(file_path.name.replace("（5月产品监督档案）", "（4月产品监督档案）")),
                actions,
                blockers,
                apply,
                may_dir,
                "may_product_archive_month_fix",
            )
    archive_office_locks(may_dir / "4月巡查", actions, blockers, apply)


def organize_template_dir(template_dir, actions, blockers, apply):
    template_dir = Path(template_dir)
    for old_name, new_name in TEMPLATE_RENAMES.items():
        move_or_rename(template_dir / old_name, template_dir / new_name, actions, blockers, apply, template_dir, "template_rename")
    archive_office_locks(template_dir, actions, blockers, apply)


def supplement_template_dir_from_may(may_dir, template_dir, actions, blockers, apply):
    may_dir = Path(may_dir)
    template_dir = Path(template_dir)
    for source_names, target_name in MAY_TEMPLATE_SOURCES.items():
        source = next((may_dir / name for name in source_names if (may_dir / name).exists()), may_dir / source_names[0])
        copy_if_missing(
            source,
            template_dir / target_name,
            actions,
            blockers,
            apply,
            DEFAULT_WORK_ROOT,
            "supplement_template_from_may",
        )


def archive_june_template_copies(june_month_dir, template_dir, actions, blockers, apply):
    june_month_dir = Path(june_month_dir)
    archive_dir = june_month_dir / "_模板副本归档"
    known_hashes = collect_known_template_hashes(template_dir)
    for name in sorted(JUNE_TEMPLATE_COPY_NAMES):
        file_path = june_month_dir / name
        if not file_path.exists() or not file_path.is_file():
            continue
        digest = sha256_file(file_path)
        if digest not in known_hashes:
            add_blocker(blockers, "疑似模板副本哈希不在模板库中，未归档", path=file_path, sha256=digest)
            continue
        move_or_rename(file_path, archive_dir / file_path.name, actions, blockers, apply, june_month_dir, "archive_june_template_copy")


def clear_docx_paragraph_colors(path, paragraphs):
    document = Document(str(path))
    changed = []
    for paragraph_no in paragraphs:
        if paragraph_no < 1 or paragraph_no > len(document.paragraphs):
            continue
        paragraph = document.paragraphs[paragraph_no - 1]
        for run in paragraph.runs:
            r_pr = run._element.rPr
            if r_pr is None:
                continue
            color = r_pr.find(qn("w:color"))
            if color is not None:
                r_pr.remove(color)
        changed.append({"paragraph": paragraph_no, "text": paragraph.text})
    document.save(str(path))
    return changed


def clear_excel_cell_colors(path, cells):
    import win32com.client

    excel = win32com.client.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    marked = []
    try:
        workbook = excel.Workbooks.Open(str(Path(path).resolve()))
        try:
            for cell_spec in cells:
                sheet = workbook.Worksheets(cell_spec["sheet"])
                cell = sheet.Cells(cell_spec["row"], cell_spec["col"])
                cell.Font.ColorIndex = -4105
                marked.append(
                    {
                        "path": str(path),
                        "sheet": cell_spec["sheet"],
                        "row": cell_spec["row"],
                        "col": cell_spec["col"],
                        "value": str(cell.Value),
                    }
                )
            workbook.Save()
        finally:
            workbook.Close(SaveChanges=True)
    finally:
        excel.Quit()
    return marked


def organize_june_dir(june_month_dir, actions, blockers, apply, template_dir):
    june_month_dir = Path(june_month_dir)
    archive_june_template_copies(june_month_dir, template_dir, actions, blockers, apply)

    product = june_month_dir / PRODUCT_REGISTER_NORMAL_NAME
    for old_name in [
        f"{PENDING_PREFIX}（5月）产品巡查底册（不发）.docx",
        "（5月）产品巡查底册（不发）.docx",
        "（5月）产品巡查底册.docx",
    ]:
        move_or_rename(june_month_dir / old_name, product, actions, blockers, apply, june_month_dir, "june_product_register_rename")

    network_dir = june_month_dir / "2026年5月联网监测基础信息考评明细表"
    old_base_info = network_dir / "5月基础信息考评截图.xls"
    base_info = june_month_dir / BASE_INFO_NORMAL_NAME
    for old_path in [
        june_month_dir / f"{PENDING_PREFIX}5月基础信息考评截图（不发）.xls",
        june_month_dir / "5月基础信息考评截图（不发）.xls",
        old_base_info,
    ]:
        move_or_rename(old_path, base_info, actions, blockers, apply, june_month_dir, "june_base_info_move")

    if apply and not blockers:
        if product.exists():
            marked = clear_docx_paragraph_colors(product, PRODUCT_REGISTER_PARAGRAPHS)
            actions.append({"kind": "clear_docx_misapplied_red", "status": "done", "path": str(product), "cleared": marked})
        else:
            add_blocker(blockers, "产品巡查底册不存在，无法撤销误标红", path=product)

        excel_marks = []
        if base_info.exists():
            excel_marks.extend(clear_excel_cell_colors(base_info, [{"sheet": "宜兴", "row": 4, "col": 6}]))
        else:
            add_blocker(blockers, "基础信息考评截图不存在，无法撤销误标红", path=base_info)
        network_stats = network_dir / "联网监测统计表.xls"
        if network_stats.exists():
            excel_marks.extend(clear_excel_cell_colors(network_stats, [{"sheet": "Sheet1", "row": 7, "col": 12}]))
        else:
            add_blocker(blockers, "联网监测统计表不存在，无法撤销误标红", path=network_stats)
        if excel_marks:
            actions.append({"kind": "clear_excel_misapplied_red", "status": "done", "cleared": excel_marks})
    else:
        actions.append(
            {
                "kind": "clear_misapplied_red",
                "status": "planned",
                "docx": str(product),
                "docx_paragraphs": PRODUCT_REGISTER_PARAGRAPHS,
                "excel_cells": [
                    {"path": str(base_info), "sheet": "宜兴", "row": 4, "col": 6},
                    {"path": str(network_dir / "联网监测统计表.xls"), "sheet": "Sheet1", "row": 7, "col": 12},
                ],
            }
        )


def run(args):
    actions = []
    blockers = []
    apply = args.apply

    organize_template_model(args.template_dir, actions, blockers, apply)
    instantiate_bulletin_dir(
        args.bulletin_dir,
        args.bulletin_year,
        args.bulletin_month,
        args.score_year,
        args.score_month,
        args.template_dir,
        actions,
        blockers,
        apply,
    )

    return {
        "mode": "apply" if apply else "dry-run",
        "bulletin": {
            "dir": str(Path(args.bulletin_dir)),
            "year": args.bulletin_year,
            "month": args.bulletin_month,
        },
        "score": {
            "year": args.score_year,
            "month": args.score_month,
            "dir": str(Path(args.bulletin_dir) / f"{args.score_month}月巡查"),
        },
        "template_dir": str(Path(args.template_dir)),
        "actions": actions,
        "blockers": blockers,
        "ok": not blockers,
    }


def main():
    parser = argparse.ArgumentParser(description="按通报月份目录模型整理月度产品成绩登记目录。")
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--dry-run", action="store_true", help="只预览动作，不修改文件")
    mode.add_argument("--apply", action="store_true", help="执行安全重命名、移动、复制和指定误生成文件删除")
    parser.add_argument("--bulletin-dir", default=str(DEFAULT_BULLETIN_DIR), help="通报月份目录，例如 ...\\26年\\6月通报")
    parser.add_argument("--bulletin-year", type=int, default=2026, help="通报目录年份")
    parser.add_argument("--bulletin-month", type=int, default=6, help="通报目录月份")
    parser.add_argument("--score-year", type=int, default=2026, help="巡查/成绩所属年份")
    parser.add_argument("--score-month", type=int, default=5, help="巡查/成绩所属月份")
    parser.add_argument("--template-dir", default=str(DEFAULT_TEMPLATE_DIR))
    args = parser.parse_args()
    result = run(args)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if result["blockers"]:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
