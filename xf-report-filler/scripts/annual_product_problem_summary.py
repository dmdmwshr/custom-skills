import argparse
import fnmatch
import json
import os
import re
import shutil
import sys
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
CONFIG_PATH = SKILL_DIR / "resources" / "annual_problem_summary.json"
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import monthly_grade_register as grade_register

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")


MONTH_FOLDER_RE = re.compile(r"(?<!\d)(\d{1,2})жңҲе·ЎжҹҘ")
MONTH_RE = re.compile(r"(?<!\d)(\d{1,2})жңҲ")
PROJECT_CODE_RE = re.compile(r"^([0-9A-Z]{10,})\s*(.*)$")
LEGACY_BRIGADE_NAMES = ["жұҹйҳҙ", "е®ңе…ҙ", "жўҒжәӘ", "й”Ўеұұ", "жғ еұұ", "ж»Ёж№–", "ж–°еҗҙ", "з»ҸејҖ"]
LEGACY_BRIGADE_SET = set(LEGACY_BRIGADE_NAMES + [f"{name}еӨ§йҳҹ" for name in LEGACY_BRIGADE_NAMES])


def load_config(path=CONFIG_PATH):
    return json.loads(Path(path).read_text(encoding="utf-8"))


def utc_now_text():
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def infer_month(path):
    path = Path(path)
    for part in reversed(path.parts[:-1]):
        match = MONTH_FOLDER_RE.search(part)
        if match:
            return int(match.group(1)), "folder"
    match = MONTH_RE.search(path.name)
    if match:
        return int(match.group(1)), "filename"
    return None, "unresolved"


def is_modified_version(path, config):
    name = Path(path).name
    return any(marker in name for marker in config["source_discovery"]["modified_markers"])


def is_excluded_source(path, config):
    name = Path(path).name
    if name.startswith(tuple(config["source_discovery"].get("ignore_prefixes", []))):
        return True
    return any(fnmatch.fnmatch(name, pattern) for pattern in config["source_discovery"].get("exclude_patterns", []))


def discover_registers(year_root, config):
    year_root = Path(year_root)
    by_month = defaultdict(list)
    blockers = []
    source_types = config["source_discovery"].get("source_types") or [
        {
            "type": "product_register_docx",
            "label": "ж–°зүҲдә§е“Ғе·ЎжҹҘеә•еҶҢ",
            "pattern": config["source_discovery"]["pattern"],
            "priority": 1,
        }
    ]
    seen = set()
    for source_type in source_types:
        for path in sorted(year_root.glob(source_type["pattern"])):
            if path in seen or is_excluded_source(path, config):
                continue
            seen.add(path)
            month, source = infer_month(path)
            item = {
                "path": str(path),
                "name": path.name,
                "month": month,
                "month_source": source,
                "source_type": source_type["type"],
                "source_label": source_type.get("label", source_type["type"]),
                "source_priority": int(source_type.get("priority", 99)),
                "modified_version": is_modified_version(path, config),
                "mtime": path.stat().st_mtime,
            }
            if not month or month < 1 or month > 12:
                blockers.append({"message": "ж— жі•иҜҶеҲ«дә§е“Ғе·ЎжҹҘеә•еҶҢжңҲд»Ҫ", **item})
                continue
            by_month[month].append(item)
    return by_month, blockers


def select_registers(by_month):
    selected = []
    skipped = []
    warnings = []
    for month in sorted(by_month):
        candidates = by_month[month]
        ranked = sorted(
            candidates,
            key=lambda item: (
                item.get("source_priority", 99),
                0 if item["modified_version"] else 1,
                -item["mtime"],
                item["name"],
            ),
        )
        chosen = ranked[0]
        selected.append(chosen)
        for item in ranked[1:]:
            skipped_item = dict(item)
            skipped_item["reason"] = f"{month}жңҲеӯҳеңЁеӨҡд»Ҫе№ҙеәҰжұҮжҖ»жәҗпјҢе·ІжҢүжәҗзұ»еһӢе’Ңдҝ®ж”№зүҲдјҳе…Ҳи§„еҲҷйҮҮз”Ё {chosen['name']}"
            skipped.append(skipped_item)
        if len(ranked) > 1:
            warnings.append(
                {
                    "type": "duplicate_month_register",
                    "message": f"{month}жңҲеӯҳеңЁеӨҡд»Ҫе№ҙеәҰжұҮжҖ»жәҗпјҢе·ІйҮҮз”Ё {chosen['name']}пјҢе…¶дҪҷеҲ—дёәжңӘйҮҮз”ЁгҖӮ",
                    "month": month,
                    "selected": chosen["path"],
                    "skipped": [item["path"] for item in ranked[1:]],
                }
            )
    return selected, skipped, warnings


def paragraph_has_yellow(paragraph):
    return grade_register.paragraph_has_yellow_highlight(paragraph)


def document_from_source(path):
    path = Path(path)
    try:
        return Document(str(path)), None
    except Exception:
        if path.suffix.lower() != ".doc":
            raise

    try:
        import win32com.client
    except Exception as exc:
        raise RuntimeError(f"ж— жі•еҠ иҪҪ Word COMпјҢдёҚиғҪиҜ»еҸ–ж—§ .docпјҡ{path}") from exc

    temp_path = Path(os.environ.get("TEMP", str(path.parent))) / f"xf-annual-source-{os.getpid()}-{datetime.now().strftime('%H%M%S%f')}.docx"
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    doc = None
    try:
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        doc.SaveAs2(str(temp_path.resolve()), FileFormat=16)
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()
    return Document(str(temp_path)), temp_path


def document_lines(path):
    temp_path = None
    try:
        document, temp_path = document_from_source(path)
        return [
            {
                "paragraph": index,
                "text": paragraph.text.strip(),
                "yellow": paragraph_has_yellow(paragraph),
            }
            for index, paragraph in enumerate(document.paragraphs, 1)
            if paragraph.text.strip()
        ]
    finally:
        if temp_path:
            temp_path.unlink(missing_ok=True)


def parse_register_entries(path, month, config):
    lines = document_lines(path)
    blocks = []
    current = None
    for item in lines:
        text = item["text"]
        if text.startswith("еӨ§йҳҹпјҡ"):
            if current:
                blocks.append(current)
            current = {"brigade": text.split("пјҡ", 1)[1].strip(), "lines": []}
        elif current:
            current["lines"].append(item)
    if current:
        blocks.append(current)

    entries = []
    blockers = []
    warnings = []
    if not blocks:
        blockers.append({"message": "дә§е“Ғе·ЎжҹҘеә•еҶҢжңӘиҜҶеҲ«еҲ°вҖңеӨ§йҳҹпјҡвҖқеҲҶеқ—", "path": str(path), "month": month})
        return entries, warnings, blockers

    for block in blocks:
        brigade = block["brigade"]
        case = {"title": "", "code": "", "filer": "", "inspector": ""}
        pending = None
        for item in block["lines"]:
            text = item["text"]
            paragraph = item["paragraph"]
            yellow = bool(item.get("yellow"))
            if text.startswith("йўҳеҗҚпјҡ"):
                case["title"] = text.split("пјҡ", 1)[1].strip()
                continue
            if text.startswith("зј–еҸ·пјҡ"):
                case["code"] = text.split("пјҡ", 1)[1].strip()
                continue
            if text.startswith("з«ӢеҚ·дәәпјҡ"):
                case["filer"] = text.split("пјҡ", 1)[1].strip()
                continue
            if text.startswith("жЈҖжҹҘдәәпјҡ"):
                case["inspector"] = text.split("пјҡ", 1)[1].strip()
                continue
            if text == "й”ҷиҜҜ" or re.fullmatch(r"\d+\s*[гҖҒ.пјҺ]?", text):
                continue
            if text.startswith("жүЈеҲҶ"):
                if pending:
                    append_entry(entries, warnings, path, month, brigade, case, pending, text, paragraph, yellow, config)
                    pending = None
                continue
            cleaned = grade_register.clean_error_line(text)
            if cleaned:
                if pending and not grade_register.is_blank_product_template_item(pending["raw"]):
                    append_entry(entries, warnings, path, month, brigade, case, pending, "", None, False, config)
                pending = {"raw": cleaned, "paragraph": paragraph, "yellow": yellow}
        if pending and not grade_register.is_blank_product_template_item(pending["raw"]):
            append_entry(entries, warnings, path, month, brigade, case, pending, "", None, False, config)

    if not entries:
        warnings.append({"message": "дә§е“Ғе·ЎжҹҘеә•еҶҢжңӘи§ЈжһҗеҲ°жңүж•Ҳй—®йўҳ", "path": str(path), "month": month})
    return entries, warnings, blockers


def normalize_legacy_brigade(text):
    text = text.strip()
    if text.endswith("еӨ§йҳҹ"):
        return text
    return f"{text}еӨ§йҳҹ"


def is_legacy_brigade_line(text):
    return text.strip() in LEGACY_BRIGADE_SET


def legacy_case_from_line(text):
    match = PROJECT_CODE_RE.match(text.strip())
    if not match:
        return None
    return {
        "code": match.group(1).strip(),
        "case_info": text.strip(),
        "title": "",
        "filer": "",
        "inspector": "",
    }


def is_legacy_score_summary(text):
    compact = grade_register.norm_text(text)
    if not compact:
        return True
    if re.fullmatch(r"жүЈ\d+(?:\.\d+)?(?:еҲҶ)?", compact):
        return True
    if re.search(r"(еҠ иө·жқҘ|еҗҲи®Ў|е…ұи®Ў|е…ұ).*жүЈ.*\d", compact):
        return True
    return False


def is_score_or_clause_only(text):
    text = str(text or "").strip()
    compact = grade_register.norm_text(text)
    if not compact:
        return True
    patterns = [
        r"^\d+(?:\.\d+)?еҲҶ$",
        r"^\d+\s+\d+(?:\.\d+)?\s*еҲҶ$",
        r"^\d+\s*[гҖҒ.пјҺ]\s*\d+(?:\.\d+)?\s*еҲҶ$",
        r"^3\s*[.пјҺ]\s*\d+\s*[a-zA-Z]?\s*\d+(?:\.\d+)?\s*еҲҶ$",
    ]
    return any(re.fullmatch(pattern, text) or re.fullmatch(pattern, compact) for pattern in patterns)


def is_standalone_ps(text):
    compact = grade_register.norm_text(text).lower().replace("(", "пјҲ").replace(")", "пјү")
    return compact.startswith("пјҲpsпјҡ") or compact.startswith("пјҲps:")


def clean_legacy_problem_text(text):
    cleaned = grade_register.clean_error_line(text)
    cleaned = re.sub(r"[пјҢ,пјӣ;гҖҒ]?\s*\d+(?:\.\d+)?\s*[пјҢ,гҖҒ]?\s*жүЈ\s*\d+(?:\.\d+)?\s*(?:еҲҶ)?\s*$", "", cleaned)
    cleaned = re.sub(r"[пјҢ,пјӣ;гҖҒ]?\s*жүЈ\s*\d+(?:\.\d+)?\s*(?:еҲҶ)?\s*$", "", cleaned)
    return cleaned.strip(" пјӣ;пјҢ,гҖҒ")


def parse_legacy_online_patrol_entries(path, month, config):
    lines = document_lines(path)
    entries = []
    warnings = []
    blockers = []
    current_brigade = None
    current_case = None
    last_entry = None

    for item in lines:
        text = item["text"].strip()
        if is_legacy_brigade_line(text):
            current_brigade = normalize_legacy_brigade(text)
            current_case = None
            last_entry = None
            continue
        case = legacy_case_from_line(text)
        if case:
            current_case = case
            last_entry = None
            continue
        if not current_brigade:
            continue
        if is_standalone_ps(text):
            if last_entry:
                last_entry.setdefault("legacy_notes", []).append(text)
            else:
                warnings.append(
                    {
                        "type": "orphan_legacy_note",
                        "message": "ж—§зүҲзҪ‘дёҠе·ЎжҹҘжәҗеӯҳеңЁж— жі•еҪ’еұһзҡ„ ps еӨҮжіЁгҖӮ",
                        "path": str(path),
                        "month": int(month),
                        "paragraph": item["paragraph"],
                        "note": text,
                    }
                )
            continue
        if is_legacy_score_summary(text):
            continue
        cleaned = clean_legacy_problem_text(text)
        if not cleaned:
            continue
        public_description = grade_register.public_product_description(cleaned)
        entry = {
            "source_path": str(path),
            "source_type": "legacy_online_patrol_doc",
            "month": int(month),
            "brigade": current_brigade,
            "brigade_short": grade_register.brigade_short(current_brigade),
            "title": "",
            "code": current_case.get("code", "") if current_case else "",
            "case_info": current_case.get("case_info", "") if current_case else "",
            "filer": "",
            "inspector": "",
            "raw_error": text,
            "public_description": public_description,
            "yellow": bool(item.get("yellow")),
            "deduction_line": "",
            "deduction_value": None,
            "general_index": None,
            "paragraph": item["paragraph"],
            "deduction_paragraph": None,
            "review_required": True,
            "review_notes": ["ж—§зүҲзҪ‘дёҠе·ЎжҹҘжәҗпјҢйңҖдәәе·ҘеӨҚж ёгҖӮ"],
            "legacy_notes": [],
        }
        if not current_case:
            entry["review_notes"].append("жңӘеҢ№й…ҚеҲ°жЎҲеҚ·дҝЎжҒҜиЎҢгҖӮ")
        if entry["yellow"]:
            entry["review_notes"].append(config["yellow_problem_policy"]["review_note"])
            warnings.append(
                {
                    "type": "yellow_problem",
                    "message": config["yellow_problem_policy"]["review_note"],
                    "path": str(path),
                    "month": int(month),
                    "brigade": current_brigade,
                    "paragraph": item["paragraph"],
                    "raw_error": text,
                }
            )
        warnings.append(
            {
                "type": "legacy_online_patrol_problem",
                "message": "ж—§зүҲзҪ‘дёҠе·ЎжҹҘжәҗй—®йўҳжқЎзӣ®е·Ізәіе…Ҙе№ҙеәҰжұҮжҖ»пјҢйңҖдәәе·ҘеӨҚж ёгҖӮ",
                "path": str(path),
                "month": int(month),
                "brigade": current_brigade,
                "paragraph": item["paragraph"],
                "raw_error": text,
            }
        )
        entries.append(entry)
        last_entry = entry

    if not entries:
        blockers.append({"message": "ж—§зүҲдә§е“Ғзӣ‘зқЈзҪ‘дёҠе·ЎжҹҘжәҗжңӘи§ЈжһҗеҲ°жңүж•Ҳй—®йўҳ", "path": str(path), "month": month})
    return entries, warnings, blockers


def append_entry(entries, warnings, path, month, brigade, case, pending, deduction_line, deduction_paragraph, deduction_yellow, config):
    raw = pending["raw"]
    if grade_register.is_blank_product_template_item(raw, deduction_line) or is_score_or_clause_only(raw):
        return
    is_yellow = bool(pending.get("yellow") or deduction_yellow)
    public_description = grade_register.public_product_description(raw)
    value = grade_register.parse_deduction_value(deduction_line) if deduction_line else None
    general_index = grade_register.parse_general_index(deduction_line) if deduction_line else None
    review_required = bool(is_yellow or not deduction_line or value is None or general_index is None)
    review_notes = []
    if is_yellow:
        review_notes.append(config["yellow_problem_policy"]["review_note"])
    if not deduction_line:
        review_notes.append("жңӘеҢ№й…ҚеҲ°жүЈеҲҶиЎҢпјҢйңҖдәәе·ҘеӨҚж ёгҖӮ")
    elif value is None or general_index is None:
        review_notes.append("жүЈеҲҶиЎҢзјәе°‘еҸҜи§ЈжһҗжқЎж¬ҫжҲ–еҲҶеҖјпјҢйңҖдәәе·ҘеӨҚж ёгҖӮ")
    entry = {
        "source_path": str(path),
        "source_type": "product_register_docx",
        "month": int(month),
        "brigade": brigade,
        "brigade_short": grade_register.brigade_short(brigade),
        "title": case.get("title", ""),
        "code": case.get("code", ""),
        "filer": case.get("filer", ""),
        "inspector": case.get("inspector", ""),
        "raw_error": raw,
        "public_description": public_description,
        "yellow": is_yellow,
        "deduction_line": deduction_line,
        "deduction_value": value,
        "general_index": general_index,
        "paragraph": pending.get("paragraph"),
        "deduction_paragraph": deduction_paragraph,
        "review_required": review_required,
        "review_notes": review_notes,
    }
    if is_yellow:
        warnings.append(
            {
                "type": "yellow_problem",
                "message": config["yellow_problem_policy"]["review_note"],
                "path": str(path),
                "month": int(month),
                "brigade": brigade,
                "paragraph": pending.get("paragraph"),
                "raw_error": raw,
            }
        )
    if review_required and not is_yellow:
        warnings.append(
            {
                "type": "low_confidence_problem",
                "message": "е№ҙеәҰжұҮжҖ»й—®йўҳжқЎзӣ®йңҖиҰҒдәәе·ҘеӨҚж ёгҖӮ",
                "path": str(path),
                "month": int(month),
                "brigade": brigade,
                "paragraph": pending.get("paragraph"),
                "raw_error": raw,
                "review_notes": review_notes,
            }
        )
    entries.append(entry)


def brigade_sort_key(brigade, config):
    order = [grade_register.brigade_short(item) for item in config["brigade_order"]]
    short = grade_register.brigade_short(brigade)
    if short in order:
        return (order.index(short), short)
    return (len(order), short)


def grouped_entries(entries, config):
    brigades = defaultdict(list)
    for entry in entries:
        brigades[entry["brigade"]].append(entry)
    return [(brigade, brigades[brigade]) for brigade in sorted(brigades, key=lambda value: brigade_sort_key(value, config))]


def set_run_font(run, font_name, size_pt=None):
    run.font.name = font_name
    if size_pt:
        run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, font_name, size_pt=None):
    for run in paragraph.runs:
        set_run_font(run, font_name, size_pt)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def case_label(entry):
    parts = []
    if entry.get("code"):
        parts.append(f"зј–еҸ·пјҡ{entry['code']}")
    if entry.get("case_info"):
        parts.append(f"жЎҲеҚ·дҝЎжҒҜпјҡ{entry['case_info']}")
    if entry.get("title"):
        parts.append(f"йўҳеҗҚпјҡ{entry['title']}")
    if entry.get("filer"):
        parts.append(f"з«ӢеҚ·дәәпјҡ{entry['filer']}")
    if entry.get("review_required"):
        parts.append("йңҖдәәе·ҘеӨҚж ё")
    return "пјӣ".join(parts) if parts else "жңӘи®°еҪ•жЎҲеҚ·дҝЎжҒҜ"


def add_issue_group(doc, description, items, config):
    style = config["word_style"]
    paragraph = doc.add_paragraph(style=None)
    run = paragraph.add_run(description)
    set_run_font(run, style["body_font"], style["body_size_pt"])
    if any(item.get("yellow") for item in items):
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for entry in items:
        case_paragraph = doc.add_paragraph(style=None)
        case_paragraph.paragraph_format.left_indent = Pt(24)
        case_run = case_paragraph.add_run(case_label(entry))
        set_run_font(case_run, style.get("fallback_body_font") or style["body_font"], style["case_size_pt"])


def add_grouped_content(doc, entries, config):
    first_brigade = True
    for brigade, brigade_entries in grouped_entries(entries, config):
        if not first_brigade:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        first_brigade = False
        heading = doc.add_heading(brigade, level=1)
        set_paragraph_font(heading, config["word_style"]["heading_font"], 16)
        by_month = defaultdict(list)
        for entry in brigade_entries:
            by_month[entry["month"]].append(entry)
        for month in sorted(by_month):
            month_heading = doc.add_heading(f"{month}жңҲ", level=2)
            set_paragraph_font(month_heading, config["word_style"]["heading_font"], 15)
            by_description = defaultdict(list)
            for entry in by_month[month]:
                by_description[entry["public_description"]].append(entry)
            for description in sorted(by_description):
                add_issue_group(doc, description, by_description[description], config)


def build_document(entries, year, config):
    doc = Document()
    style = config["word_style"]
    title_text = style["title"].format(year=year)
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(title_text)
    title_run.bold = True
    set_run_font(title_run, style["title_font"], 22)

    toc_title = doc.add_paragraph()
    toc_title_run = toc_title.add_run(style["toc_title"])
    toc_title_run.bold = True
    set_run_font(toc_title_run, style["heading_font"], 16)
    toc_paragraph = doc.add_paragraph()
    add_toc(toc_paragraph)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_grouped_content(doc, entries, config)
    return doc


def update_toc_with_word(path):
    try:
        import win32com.client
    except Exception as exc:
        raise RuntimeError(f"ж— жі•еҠ иҪҪ Word COMпјҢдёҚиғҪжӣҙж–°зӣ®еҪ•пјҡ{exc}") from exc

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    doc = None
    try:
        doc = word.Documents.Open(str(Path(path).resolve()), ReadOnly=False)
        if doc.TablesOfContents.Count < 1:
            raise RuntimeError("е№ҙеәҰжұҮжҖ»ж–ҮжЎЈжңӘз”ҹжҲҗзӣ®еҪ•еӯ—ж®ө")
        doc.TablesOfContents(1).Update()
        doc.Save()
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


def write_annual_doc(entries, year, output_path, config, force=False, update_toc=True):
    output_path = Path(output_path)
    if output_path.exists() and not force:
        raise FileExistsError(f"зӣ®ж Үж–Үд»¶е·ІеӯҳеңЁпјҢдҪҝз”Ё --force иҰҶзӣ–пјҡ{output_path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = output_path.parent / f".annual-problem-summary-{os.getpid()}-{datetime.now().strftime('%H%M%S%f')}.docx"
    try:
        doc = build_document(entries, year, config)
        doc.save(temp_path)
        if update_toc:
            update_toc_with_word(temp_path)
        if output_path.exists():
            output_path.unlink()
        shutil.move(str(temp_path), str(output_path))
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise
    return output_path


def output_path_for(year_root, year, output, config):
    if output:
        return Path(output)
    return Path(year_root) / config["output_file"].format(year=year)


def build_review(year_root, year, config):
    by_month, discovery_blockers = discover_registers(year_root, config)
    selected, skipped, warnings = select_registers(by_month)
    blockers = list(discovery_blockers)
    if not selected:
        blockers.append({"message": "е№ҙеәҰж №зӣ®еҪ•жңӘжүҫеҲ°еҸҜйҮҮз”Ёзҡ„дә§е“Ғе·ЎжҹҘеә•еҶҢ", "year_root": str(Path(year_root))})
    entries = []
    for source in selected:
        if source.get("source_type") == "legacy_online_patrol_doc":
            source_entries, source_warnings, source_blockers = parse_legacy_online_patrol_entries(source["path"], source["month"], config)
        else:
            source_entries, source_warnings, source_blockers = parse_register_entries(source["path"], source["month"], config)
        entries.extend(source_entries)
        warnings.extend(source_warnings)
        blockers.extend(source_blockers)
    known_order = {grade_register.brigade_short(item) for item in config["brigade_order"]}
    for entry in entries:
        if entry["brigade_short"] not in known_order:
            warnings.append(
                {
                    "type": "unknown_brigade",
                    "message": "е№ҙеәҰжұҮжҖ»еҸ‘зҺ°жңӘзҹҘеӨ§йҳҹпјҢе·ІжҺ’еңЁж–Үжң«гҖӮ",
                    "brigade": entry["brigade"],
                    "source_path": entry["source_path"],
                    "month": entry["month"],
                }
            )
    return {
        "ok": not blockers,
        "generated_at": utc_now_text(),
        "year_root": str(Path(year_root)),
        "year": int(year),
        "selected_sources": selected,
        "skipped_sources": skipped,
        "entries": entries,
        "warnings": warnings,
        "blockers": blockers,
    }


def write_review_json(path, review):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(review, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def main():
    config = load_config()
    parser = argparse.ArgumentParser(
        description="з”ҹжҲҗжҲ–е®Ўи®Ўе№ҙеәҰдә§е“Ғзӣ‘зқЈеә•еҶҢй—®йўҳжұҮжҖ»гҖӮ",
        epilog="еҸӮиҖғпјҡreferences/annual_problem_summary/00_workflow_router.mdгҖӮ",
    )
    parser.add_argument("--year-root", default=config["default_year_root"], help="е№ҙеәҰж №зӣ®еҪ•пјҢдҫӢеҰӮ ...\\26е№ҙ")
    parser.add_argument("--year", type=int, required=True)
    parser.add_argument("--output", help="иҫ“еҮә .docx и·Ҝеҫ„пјӣй»ҳи®ӨеҶҷе…Ҙе№ҙеәҰж №зӣ®еҪ•")
    parser.add_argument("--review-json", help="еҶҷеҮәйҖҗжқЎе®Ўи®Ў JSON")
    parser.add_argument("--force", action="store_true", help="иҰҶзӣ–е·ІеӯҳеңЁиҫ“еҮәж–Үд»¶")
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--dry-run", action="store_true", help="еҸӘе®Ўи®ЎпјҢдёҚз”ҹжҲҗ Word")
    mode.add_argument("--apply", action="store_true", help="з”ҹжҲҗ Word жұҮжҖ»")
    args = parser.parse_args()

    review = build_review(args.year_root, args.year, config)
    output_path = output_path_for(args.year_root, args.year, args.output, config)
    review["output_path"] = str(output_path)
    if args.review_json:
        write_review_json(args.review_json, review)
    print(json.dumps(review, ensure_ascii=False, indent=2))
    if review["blockers"]:
        raise SystemExit(2)
    if args.apply:
        write_annual_doc(review["entries"], args.year, output_path, config, force=args.force, update_toc=True)
        print(json.dumps({"generated": str(output_path)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
