import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import annual_product_problem_summary as annual


def write_register(path, brigade="宜兴大队", error="卷内应有的某文书缺失（缺少审批表）", yellow=False, deduction=True):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph(f"大队：{brigade}")
    doc.add_paragraph("题名：测试案卷")
    doc.add_paragraph("编号：WX-001")
    doc.add_paragraph("立卷人：张三")
    doc.add_paragraph("错误")
    issue = doc.add_paragraph()
    run = issue.add_run(f"1、{error}")
    if yellow:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if deduction:
        doc.add_paragraph("扣分：3.13 a 0.5分")
    doc.save(path)
    return path


def write_no_brigade_register(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("题名：测试案卷")
    doc.add_paragraph("1、卷内应有的某文书缺失")
    doc.save(path)
    return path


def write_legacy_online_patrol(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("江阴")
    doc.add_paragraph("32002221C202500062江阴安居酒店管理有限公司缪晓清(主)、过超")
    doc.add_paragraph("未上传初次检查产品证书查询截图（ps：只有佛山四甲，没有广州锐盾）")
    doc.add_paragraph("消防产品质量监督抽查抽样单抽样人员未签字（未看见纸质版，无法核实签字）0.2，扣0.1")
    doc.add_paragraph("未上传广州锐盾的不合格检测报告（有电子版，未上传纸质版扣0.5分吗）")
    doc.add_paragraph("宜兴")
    doc.add_paragraph("32002222C202500058宜兴市宜城街道悦昕餐厅魏青凯(主)、周智泰")
    doc.add_paragraph("证据保全决定书填写不规范（ps：因调查“xxx单位使用不合格消防产品一案”）")
    doc.add_paragraph("证据保全清单签名与当事人名字不符（ps：错别字）")
    doc.add_paragraph("这个两项加起来扣除了0.5")
    doc.add_paragraph("经开")
    doc.add_paragraph("32113513C202500026经开区雀隐庐茶艺棋牌店（个体工商户）巫超(主)、张杰")
    doc.add_paragraph("证据保全决定书格式错误")
    doc.add_paragraph("（ps：整改照片未开盒拍摄新面罩的外观情况）")
    doc.save(path)
    return path


class AnnualProductProblemSummaryTests(unittest.TestCase):
    def setUp(self):
        self.config = annual.load_config()

    def test_infers_month_from_score_folder_and_prefers_modified_version(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            old_path = write_register(root / "4月通报" / "3月巡查" / "0、（3月不发）产品巡查底册.docx")
            modified_path = write_register(root / "4月通报" / "3月巡查" / "0、（3月不发）（修改版本）产品巡查底册.docx")
            old_path.touch()
            modified_path.touch()

            review = annual.build_review(root, 2026, self.config)

            self.assertTrue(review["ok"])
            self.assertEqual(len(review["selected_sources"]), 1)
            self.assertEqual(Path(review["selected_sources"][0]["path"]).name, modified_path.name)
            self.assertEqual(review["selected_sources"][0]["month"], 3)
            self.assertEqual(len(review["skipped_sources"]), 1)
            self.assertEqual(Path(review["skipped_sources"][0]["path"]).name, old_path.name)

    def test_discovers_legacy_online_patrol_doc_and_excludes_archive_docs(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            legacy_path = write_legacy_online_patrol(
                root / "2月通报" / "1月巡查" / "1月产品成绩" / "（不发）2026年1月产品监督网上巡查.doc"
            )
            write_legacy_online_patrol(root / "2月通报" / "1月巡查" / "1月产品成绩" / "滨湖（1月档案）.doc")

            review = annual.build_review(root, 2026, self.config)

            self.assertTrue(review["ok"])
            self.assertEqual(len(review["selected_sources"]), 1)
            self.assertEqual(Path(review["selected_sources"][0]["path"]).name, legacy_path.name)
            self.assertEqual(review["selected_sources"][0]["month"], 1)
            self.assertEqual(review["selected_sources"][0]["source_type"], "legacy_online_patrol_doc")

    def test_legacy_online_patrol_parses_problem_lines_and_keeps_ps_in_review_only(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            write_legacy_online_patrol(
                root / "2月通报" / "1月巡查" / "1月产品成绩" / "（不发）2026年1月产品监督网上巡查.doc"
            )

            review = annual.build_review(root, 2026, self.config)

            self.assertTrue(review["ok"])
            descriptions = [entry["public_description"] for entry in review["entries"]]
            self.assertIn("消防产品质量监督抽查抽样单抽样人员未签字", descriptions)
            self.assertNotIn("消防产品质量监督抽查抽样单抽样人员未签字0.2，扣0.1", descriptions)
            self.assertFalse(any("整改照片" in item["public_description"] for item in review["entries"]))
            note_entry = next(item for item in review["entries"] if item["public_description"] == "证据保全决定书格式错误")
            self.assertIn("整改照片未开盒", note_entry["legacy_notes"][0])
            self.assertTrue(note_entry["review_required"])
            self.assertEqual(note_entry["code"], "32113513C202500026")
            self.assertIn("经开区雀隐庐", note_entry["case_info"])
            self.assertTrue(any(item.get("type") == "legacy_online_patrol_problem" for item in review["warnings"]))

    def test_yellow_problem_is_included_and_public_text_strips_detail(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            write_register(
                root / "3月通报" / "2月巡查" / "2、修改版本（2月不发）产品巡查底册.docx",
                error="复查现场照片不齐全（缺少门头远景）",
                yellow=True,
            )

            review = annual.build_review(root, 2026, self.config)

            self.assertTrue(review["ok"])
            self.assertEqual(len(review["entries"]), 1)
            entry = review["entries"][0]
            self.assertTrue(entry["yellow"])
            self.assertEqual(entry["public_description"], "复查现场照片不齐全")
            self.assertIn("缺少门头远景", entry["raw_error"])
            self.assertTrue(entry["review_required"])
            self.assertTrue(any(item.get("type") == "yellow_problem" for item in review["warnings"]))

    def test_simple_early_register_without_deduction_is_low_confidence_not_blocked(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            write_register(
                root / "3月通报" / "2月巡查" / "2月产品巡查底册.docx",
                error="消防产品监督检查记录填写不规范",
                deduction=False,
            )

            review = annual.build_review(root, 2026, self.config)

            self.assertTrue(review["ok"])
            self.assertEqual(len(review["entries"]), 1)
            self.assertTrue(review["entries"][0]["review_required"])
            self.assertIn("未匹配到扣分行", "；".join(review["entries"][0]["review_notes"]))

    def test_score_only_lines_are_not_exported_as_problems(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "4月通报" / "3月巡查" / "3月产品巡查底册.docx"
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.add_paragraph("大队：锡山大队")
            doc.add_paragraph("题名：测试案卷")
            doc.add_paragraph("编号：WX-001")
            doc.add_paragraph("立卷人：张三")
            doc.add_paragraph("错误")
            doc.add_paragraph("1、无营业执照")
            doc.add_paragraph("11 0.1分")
            doc.save(path)

            review = annual.build_review(root, 2026, self.config)

            descriptions = [entry["public_description"] for entry in review["entries"]]
            self.assertIn("无营业执照", descriptions)
            self.assertNotIn("11 0.1分", descriptions)

    def test_register_without_brigade_is_blocker(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            write_no_brigade_register(root / "3月通报" / "2月巡查" / "2月产品巡查底册.docx")

            review = annual.build_review(root, 2026, self.config)

            self.assertFalse(review["ok"])
            self.assertTrue(any("未识别到“大队：”分块" in item["message"] for item in review["blockers"]))

    def test_no_registers_is_blocker(self):
        with tempfile.TemporaryDirectory() as tmp:
            review = annual.build_review(Path(tmp), 2026, self.config)

            self.assertFalse(review["ok"])
            self.assertTrue(any("未找到可采用的产品巡查底册" in item["message"] for item in review["blockers"]))

    def test_writes_docx_with_cover_navigation_headings_and_yellow_highlight(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            write_register(
                root / "3月通报" / "2月巡查" / "2月产品巡查底册.docx",
                brigade="梁溪大队",
                error="复查现场照片不齐全（缺少门头远景）",
                yellow=True,
            )
            write_register(
                root / "4月通报" / "3月巡查" / "3月产品巡查底册.docx",
                brigade="锡山大队",
                error="缺授权委托书（缺授权委托书）",
            )
            review = annual.build_review(root, 2026, self.config)
            output = root / "2026年产品监督底册问题汇总.docx"

            annual.write_annual_doc(review["entries"], 2026, output, self.config, update_toc=False)

            self.assertTrue(output.exists())
            with zipfile.ZipFile(output) as archive:
                xml = archive.read("word/document.xml").decode("utf-8")
            self.assertNotIn('TOC \\o "1-2"', xml)
            self.assertNotIn(">目录<", xml)
            self.assertIn('w:val="Heading1"', xml)
            self.assertIn('w:val="Heading2"', xml)
            self.assertIn('<w:br w:type="page"', xml)
            self.assertIn("<w:highlight", xml)
            self.assertIn("梁溪大队", xml)
            self.assertIn("锡山大队", xml)
            self.assertIn("复查现场照片不齐全", xml)
            self.assertNotIn("缺少门头远景", xml)
            self.assertNotIn("需人工复核", xml)

    def test_generated_docx_groups_issues_under_case_without_repeating_case_info_or_scores(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "3月通报" / "2月巡查" / "2月产品巡查底册.docx"
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.add_paragraph("大队：梁溪大队")
            doc.add_paragraph("题名：测试案卷")
            doc.add_paragraph("编号：WX-001")
            doc.add_paragraph("立卷人：张三")
            doc.add_paragraph("错误")
            doc.add_paragraph("1、消防产品监督检查记录填写不规范")
            doc.add_paragraph("扣分：3.13 a 0.5分")
            doc.add_paragraph("2、现场照片缺概貌（缺门头照）")
            doc.add_paragraph("扣分：3.13 b 0.2分")
            doc.save(path)
            review = annual.build_review(root, 2026, self.config)
            output = root / "2026年产品监督底册问题汇总.docx"

            annual.write_annual_doc(review["entries"], 2026, output, self.config)

            with zipfile.ZipFile(output) as archive:
                xml = archive.read("word/document.xml").decode("utf-8")
            self.assertEqual(xml.count("WX-001"), 1)
            self.assertIn("1、消防产品监督检查记录填写不规范", xml)
            self.assertIn("2、现场照片缺概貌", xml)
            self.assertNotIn("缺门头照", xml)
            self.assertNotIn("3.13", xml)
            self.assertNotIn("0.5分", xml)

    def test_generated_docx_includes_legacy_month_without_ps_or_score_tail(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            write_legacy_online_patrol(
                root / "2月通报" / "1月巡查" / "1月产品成绩" / "（不发）2026年1月产品监督网上巡查.doc"
            )
            review = annual.build_review(root, 2026, self.config)
            output = root / "2026年产品监督底册问题汇总.docx"

            annual.write_annual_doc(review["entries"], 2026, output, self.config, update_toc=False)

            with zipfile.ZipFile(output) as archive:
                xml = archive.read("word/document.xml").decode("utf-8")
            self.assertIn("1月", xml)
            self.assertIn("江阴大队", xml)
            self.assertIn("消防产品质量监督抽查抽样单抽样人员未签字", xml)
            self.assertNotIn("扣0.1", xml)
            self.assertNotIn("整改照片未开盒", xml)

    def test_review_json_can_be_written(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            write_register(root / "3月通报" / "2月巡查" / "2月产品巡查底册.docx")
            review = annual.build_review(root, 2026, self.config)
            review_path = root / "review.json"

            annual.write_review_json(review_path, review)

            loaded = json.loads(review_path.read_text(encoding="utf-8"))
            self.assertEqual(loaded["entries"][0]["raw_error"], "卷内应有的某文书缺失（缺少审批表）")


if __name__ == "__main__":
    unittest.main()
